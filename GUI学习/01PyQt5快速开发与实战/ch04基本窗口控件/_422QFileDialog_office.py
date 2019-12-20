import sys
import docx
import xlrd, xlwt
from PyQt5.QtWidgets import QApplication, QWidget, QVBoxLayout, QPushButton, QLabel, QFileDialog, QTextEdit

class FileDialogDemo(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("File Dialog demo")
        self.setGeometry(300, 300, 300, 200)
        layout = QVBoxLayout()

        self.btn1 = QPushButton("加载word")
        self.btn1.clicked.connect(self.get_word)
        layout.addWidget(self.btn1)

        self.te = QTextEdit()
        layout.addWidget(self.te)

        self.btn2 = QPushButton("新建excel")
        self.btn2.clicked.connect(self.set_excel)
        layout.addWidget(self.btn2)

        self.btn3 = QPushButton("加载excel")
        self.btn3.clicked.connect(self.get_excel)
        layout.addWidget(self.btn3)

        self.te2 = QTextEdit()
        layout.addWidget(self.te2)

        self.setLayout(layout)

    def get_word(self):
        f_name = QFileDialog.getOpenFileName(self, "open word",
                                             "E:\\", "Word Files(*.doc *.docx)")
        if f_name[0]:
            # 获取文档对象
            file = docx.Document(f_name[0])
            # 输出每一段的内容
            text = ''
            for i in range(len(file.paragraphs)):
                text += "第{}段：{}\n".format(i, file.paragraphs[i].text)
            self.te.setText(text)

    def set_excel(self):
        # 创建一个Workbook对象，这就相当于创建了一个Excel文件
        # style_compression表示是否压缩，不常用
        book = xlwt.Workbook(encoding='utf-8', style_compression=0)
        # 创建一个sheet对象，一个sheet对象对应Excel文件中的一张表格
        # cell_overwrite_ok表示是否可以覆盖单元格，默认Fasle
        sheet = book.add_sheet('test', cell_overwrite_ok=True)
        # 向表格test中添加数据
        sheet.write(0, 0, 'EnglishName')
        sheet.write(1, 0, 'daguo')
        sheet.write(0, 1, '中文名字')
        sheet.write(1, 1, '大国')
        sheet.write(0, 2, '出生日期')
        sheet.write(1, 2, '2000/1/1')
        book.save(r'.\test.xls')  # 保存以上操作

    def get_excel(self):
        f_name = QFileDialog.getOpenFileName(self, "open excel",
                                             "E:\\", "Excel Files(*.xls *.xlsx)")
        if f_name[0]:
            # 获取excel文件对象
            book = xlrd.open_workbook(f_name[0])
            sheet0 = book.sheet_by_index(0)  # 通过sheet索引获得sheet对象
            sheet_name = book.sheet_names()[0]  # 获得指定索引的sheet表名字
            sheet1 = book.sheet_by_name(sheet_name)  # 通过sheet名字来获得sheet对象
            rows = sheet0.nrows  # 获取行数
            row_data = []  # 按行获取内容
            for i in range(rows):
                row_data.append(sheet0.row_values(i))
            cols = sheet0.ncols  # 获取列数
            col_data = []  # 案列获取内容
            for i in range(cols):
                col_data.append(sheet1.col_values(i))
            # 通过坐标获取表格中的数据
            cell_value = sheet0.cell_value(0, 0)
            cell_value2 = sheet0.cell(1, 2).value
            self.te2.setText('{}\n或{}\n名字：{}，共{}行{}列\n按行输出内容：{}\n'
                             '按列输出内容：{}\n(0,0)处内容：{}\n(1,2)处内容：{}'.format(
                              sheet0, sheet1, sheet_name, rows, cols, row_data, col_data,
                              cell_value, cell_value2))


if __name__ == '__main__':
    app = QApplication(sys.argv)
    my_show = FileDialogDemo()
    my_show.show()
    sys.exit(app.exec_())
